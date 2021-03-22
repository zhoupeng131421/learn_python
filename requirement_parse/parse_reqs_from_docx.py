#!/usr/bin/python3
# _*_ coding: UTF-8 _*_

import sys
import os
import re
from docx import Document
'''
import prettytable as pt
TABLE = pt.PrettyTable()
'''

Dict = {}

def parse_table(table):
    id = ""
    Trace_from = ""
    tracefroms = []
    for row in table.rows:
        if row.cells[0].text.strip() == 'Trace from':
            Trace_from = row.cells[1].text.strip()
            tracefroms = re.split(r'\n', Trace_from)
        elif row.cells[0].text.strip() == 'Test case identifier':
            id = row.cells[1].text.strip()
    if id.strip():
        for tracefrom in tracefroms:
            if tracefrom in Dict:
                new_tracefrom = Dict[tracefrom] + '\n' + id
                Dict[tracefrom] = new_tracefrom
            else:
                Dict[tracefrom] = id

def match_doc(target_doc):
    document = Document(target_doc)
    tables = document.tables
    for table in tables:
        parse_table(table)

def main():
    global Dict
    '''
    global TABLE
    TABLE.field_names = ['Requirement', 'Proof', 'Analyses/Test cases']
    '''
    out_document = Document()

    # get target document path
    argc = len(sys.argv)
    if argc == 2:
        target_doc = os.path.abspath(sys.argv[1])
        print("target document is: ", target_doc)
    else:
        print("There no have parameter about target document, please check your input.")
        sys.exit()

    match_doc(sys.argv[1])
    table = out_document.add_table(rows = len(Dict) + 1, cols = 3, style = 'TableGrid')
    row_index = 0
    table.rows[row_index].cells[0].text = 'Requirement'
    table.rows[row_index].cells[1].text = 'Proof'
    table.rows[row_index].cells[2].text = 'Analyses/Test cases'
    row_index += 1
    for key,value in Dict.items():
        table.rows[row_index].cells[0].text = key
        table.rows[row_index].cells[1].text = 'test'
        table.rows[row_index].cells[2].text = str(value)
        row_index += 1
    out_document.save(sys.argv[1] + '.docx')

    '''
    TABLE.align = 'l'
    TABLE.padding_width = 0
    for key,value in Dict.items():
        ROW = []
        ROW.append(key)
        ROW.append('test')
        ROW.append(str(value))
        TABLE.add_row(ROW)
    print(TABLE)
    '''

if __name__ == '__main__':
    main()